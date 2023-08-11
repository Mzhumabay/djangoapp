from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from django.db import connection
import pandas as pd
import win32com.client as win32
from openpyxl import load_workbook
from io import BytesIO
import os


def get_publications_data():
    with connection.cursor() as cursor:
        cursor.execute("SELECT * FROM gdp_regions")
        rows = cursor.fetchall()
    return rows


def get_data(id):

    with connection.cursor() as cursor:
        cursor.execute(id)
        rows = cursor.fetchall()
        list_list = [list(tpl) for tpl in rows]
    return list_list


def index(request):
    links = [
    {"subsection_id": "vvp", "display_text": "Валовый внутренний продукт (ВВП)"},
    {"subsection_id": "ifo", "display_text": "Индекс физического объема (ИФО)"},
    {"subsection_id": "prod_truda", "display_text": "Производительность труда"},
    {"subsection_id": "invest_osn_kapital", "display_text": "Инвестиции в основной капитал"},
    {"subsection_id": "zanyatost_bezrabotica", "display_text": "Занятость, безработица, средняя зарплата по основным странам"},
    {"subsection_id": "potreb_ceni_proizv_ceni", "display_text": "Индекс потребительских цен и индекс цен производителей"},
    {"subsection_id": "potreb_ceni_socialno_znac_tovary", "display_text": "Индекс цен на социально-значимые потребительские товары"},
    {"subsection_id": "potreb_ceni_po_stranam", "display_text": "Индекс потребительских цен по странам"},
    {"subsection_id": "mezhdunar_rezervy_kursy_valut", "display_text": "Международные резервы и курсы валют"},
    {"subsection_id": "gosdolg_v_pp_k_vvp_po_stranam", "display_text": "Госдолг в % к ВВП по странам"},
    {"subsection_id": "kreditny_reiting", "display_text": "Кредитный рейтинг"},
    {"subsection_id": "ispolnenie_gos_bjudzheta_dohody", "display_text": "Исполнение государственного бюджета (доходы)"},
    {"subsection_id": "ispolnenie_gos_bjudzheta_zatrati_deficit", "display_text": "Исполнение государственного бюджета (затраты и дефицит)"},
    {"subsection_id": "index_pmi_sovokupnyy_po_stranam", "display_text": "Индекс PMI (совокупный) по странам"},
    {"subsection_id": "index_pmi_promyshlennost_uslugi_po_stranam", "display_text": "Индекс PMI (в промышленности и услугах) по странам"},
    {"subsection_id": "torgovy_oborot_kazahstan", "display_text": "Торговый оборот Республики Казахстан"},
    {"subsection_id": "eksport_kazahstan_po_stranam", "display_text": "Экспорт Республики Казахстан в разрезе стран"},
    {"subsection_id": "import_kazahstan_po_stranam", "display_text": "Импорт Республики Казахстан в разрезе стран"},
    {"subsection_id": "osnovnye_eksportnye_tovary", "display_text": "Основные экспортные товары"},
    {"subsection_id": "osnovnye_importnye_tovary", "display_text": "Основные импортные товары"},
    {"subsection_id": "prognoz_socialno_ekonomich_razvitie", "display_text": "Прогноз социально-экономического развития"},
    {"subsection_id": "prognoz_institut_ekonomicheskih_issledovanij", "display_text": "Прогноз Института экономических исследований"},
    {"subsection_id": "konsensus_prognoz", "display_text": "Консенсус прогноз"},
    ]
    context = {
        'links': links,
    }
    #print(links)
    return render(request, 'index.html', context)


def table_view(request):
    publications_data = get_publications_data()

    # Передаем данные в шаблон для отображения
    context = {
        'publications_data': publications_data
    }

    return render(request, 'table_view.html', context)


def excel_to_doc(request):
    # Загружаем Excel-файл с данными
    excel_file_path = 'static/doc/test_table.xlsx'
    wb = load_workbook(excel_file_path)
    ws = wb.active
# Создаем список списков для хранения данных
    data_list = []

# Проходим по строкам и столбцам таблицы и добавляем данные в список
    for row in ws.iter_rows():
        row_data = [cell.value for cell in row]
        data_list.append(row_data)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="table_data.docx"'

    doc = Document()

    # Ваш код для создания DOCX документа с данными из таблиц
    doc.add_heading('Таблица публикаций:', level=2)
    for publication in data_list:
        doc.add_paragraph(
            f"ID: {publication[0]}, Название: {publication[1]}, Название: {publication[2]}, Название: {publication[3]}")

    doc.save(response)
    return response


def download_excel(request):
    excel_file_path = 'static/doc/test_table.xlsx'  # Путь к вашему Excel файлу
    if os.path.exists(excel_file_path):
        with open(excel_file_path, 'rb') as excel_file:
            response = HttpResponse(excel_file.read(
            ), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response[
                'Content-Disposition'] = f'attachment; filename={os.path.basename(excel_file_path)}'
            return response
    else:
        # Обработка случая, если файл не найден
        return HttpResponse('Excel файл не найден', status=404)


def download_docx(request):
    publications_data = get_publications_data()
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="table_data.docx"'

    doc = Document()

    # Ваш код для создания DOCX документа с данными из таблиц
    doc.add_heading('Таблица публикаций:', level=2)
    for publication in publications_data:
        doc.add_paragraph(f"ID: {publication[0]}, Название: {publication[1]}")

    doc.save(response)
    return response


def subsection_detail(request, subsection_id):
    excel_file_path = 'static/doc/test_table.xlsx'

    # Прочитать данные из Excel файла
    df = pd.read_excel(excel_file_path)

    # Преобразовать DataFrame в HTML-таблицу
    html_table = df.to_html(index=False, classes='table table-bordered')

    return render(request, 'excel_to_html.html', {'html_table': html_table})
